#!/usr/bin/env python3
"""Generate docs/Utility-FHIR Query Validator.docx"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT = REPO_ROOT / "docs" / "Utility-FHIR Query Validator.docx"


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))


def add_numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


def add_code_block(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    shading.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shading)


def add_architecture_table(doc):
    doc.add_heading("Validation Flow (Architecture)", level=2)
    doc.add_paragraph(
        "The diagram below summarizes how a FHIR search URL is validated against "
        "the target server CapabilityStatement before any search is executed."
    )
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("1. Entry", "CLI (fhir-validate), Python API, or demo notebook"),
        ("2. Parse", "Extract resource type and query parameters from URL"),
        ("3. Load metadata", "Fetch CapabilityStatement from /metadata"),
        ("4. Validate", "Check resource type, search params, modifiers, comparators"),
        ("5. Result", "Return {valid: true/false, errors: [...]}"),
    ]
    for i, (step, detail) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = step
        c1.text = detail
        set_cell_shading(c0, "D5E8F0")
        for cell in (c0, c1):
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Mermaid equivalent:").italic = True
    add_code_block(
        doc,
        "QueryURL --> Parse --> Validate\n"
        "CapabilityStatement --> Index --> Validate\n"
        "Validate --> StaticValueSets --> Result[valid_and_errors]",
    )


def add_repo_structure_table(doc):
    doc.add_heading("Repository Structure", level=1)
    structure = [
        ("config/", "Environment template (.env.example)"),
        ("docs/", "PRD, ADR, guides, sample output"),
        ("examples/notebooks/", "Demo Jupyter notebook"),
        ("planning/", "3-week implementation plan"),
        ("scripts/", "CLI wrappers and test runners"),
        ("src/fhir_validator_agent/config/", "Settings and public server registry"),
        ("src/fhir_validator_agent/core/", "Parsing, validation rules, value sets"),
        ("src/fhir_validator_agent/infrastructure/", "HTTP: metadata fetch, OAuth"),
        ("src/fhir_validator_agent/services/", "FhirValidatorService orchestration"),
        ("tests/unit/", "Fast offline unit tests"),
        ("tests/regression/", "JSON-driven regression cases"),
        ("tests/integration/", "Live public FHIR server tests"),
    ]
    table = doc.add_table(rows=1 + len(structure), cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Path"
    hdr[1].text = "Purpose"
    set_cell_shading(hdr[0], "2E75B6")
    set_cell_shading(hdr[1], "2E75B6")
    for p in hdr[0].paragraphs + hdr[1].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
    for i, (path, purpose) in enumerate(structure, 1):
        table.rows[i].cells[0].text = path
        table.rows[i].cells[1].text = purpose


def add_out_of_scope(doc):
    doc.add_heading("Out of Scope", level=1)
    doc.add_paragraph("The utility is a pre-flight validator only. It will not:")
    sections = [
        ("Execution and retrieval", [
            "Execute searches or return FHIR resources",
            "Validate pagination, _include, chained searches, or _sort/_count",
            "Process Bundle responses",
        ]),
        ("API and agent frameworks", [
            "Expose an HTTP/REST API service",
            "Integrate with Google ADK or GenAI query generation",
            "Translate natural language to FHIR queries",
        ]),
        ("Terminology services", [
            "Look up ValueSet/CodeSystem from a terminology server",
            "Validate SNOMED, LOINC, or ICD codes at query time",
        ]),
        ("Operations", [
            "Cache CapabilityStatement with TTL policies",
            "Provide a hosted SaaS or UI dashboard",
        ]),
    ]
    for title, items in sections:
        doc.add_heading(title, level=2)
        for item in items:
            add_bullet(doc, item)


def add_testing_section(doc):
    doc.add_heading("How to Run Tests", level=1)
    doc.add_paragraph("Prerequisites: Python 3.11+ and dev dependencies.")
    add_code_block(doc, 'pip install -e ".[dev]"')
    doc.add_heading("Recommended commands", level=2)
    tests = [
        ("All offline tests (CI)", 'pytest -m "not integration"'),
        ("Unit tests only", "make test-unit"),
        ("Regression suite", "make test-regression"),
        ("Coverage report (80% gate)", "make test-cov"),
        ("Integration tests (network)", "make test-integration"),
        ("Multi-server manual script", "python3 scripts/run_all_tests.py"),
    ]
    table = doc.add_table(rows=1 + len(tests), cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Purpose"
    table.rows[0].cells[1].text = "Command"
    set_cell_shading(table.rows[0].cells[0], "D5E8F0")
    set_cell_shading(table.rows[0].cells[1], "D5E8F0")
    for i, (purpose, cmd) in enumerate(tests, 1):
        table.rows[i].cells[0].text = purpose
        table.rows[i].cells[1].text = cmd

    doc.add_paragraph()
    doc.add_paragraph(
        "Test layout: tests/unit/ (module tests), tests/regression/ (fixed JSON cases), "
        "tests/integration/ (live HAPI, Firely, Spark, WildFHIR servers)."
    )


def add_appendix_configuration(doc):
    doc.add_heading("Appendix A: Configuration", level=1)
    doc.add_heading("Quick setup", level=2)
    add_code_block(doc, "cp config/.env.example config/.env.local")
    doc.add_paragraph(
        "Edit config/.env.local. The validator loads this file via FhirValidatorService.from_env()."
    )

    doc.add_heading("Environment variables", level=2)
    env_vars = [
        ("FHIR_DEFAULT_SERVER_KEY", "hapi", "No", "Preset: hapi, firely, spark, wildfhir"),
        ("FHIR_METADATA_URL", "(from preset)", "No", "Override CapabilityStatement URL"),
        ("FHIR_SERVER_BASE", "(from preset)", "No", "Override FHIR server base URL"),
        ("FHIR_USE_AUTH", "false", "No", "Enable OAuth client-credentials"),
        ("TOKEN_URL", "—", "If auth", "OAuth token endpoint"),
        ("CLIENT_ID", "—", "If auth", "OAuth client ID"),
        ("CLIENT_SECRET", "—", "If auth", "OAuth client secret"),
    ]
    table = doc.add_table(rows=1 + len(env_vars), cols=4)
    table.style = "Table Grid"
    headers = ["Variable", "Default", "Required", "Description"]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        set_cell_shading(table.rows[0].cells[j], "2E75B6")
    for i, row in enumerate(env_vars, 1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val

    doc.add_heading("Public test servers (no authentication)", level=2)
    servers = [
        ("hapi (default)", "https://hapi.fhir.org/baseR4"),
        ("firely", "https://server.fire.ly"),
        ("spark", "https://spark.incendi.no/fhir"),
        ("wildfhir", "https://wildfhir.wildfhir.org/r4"),
    ]
    for key, url in servers:
        add_bullet(doc, f"{key}: {url}")

    doc.add_heading("OAuth example", level=2)
    add_code_block(
        doc,
        "FHIR_USE_AUTH=true\n"
        "TOKEN_URL=https://auth.example.com/oauth/token\n"
        "CLIENT_ID=your-client-id\n"
        "CLIENT_SECRET=your-client-secret",
    )

    doc.add_heading("Security practices", level=2)
    for item in [
        "Never commit config/.env.local or files with secrets",
        "Use config/.env.example as the documented template only",
        "Do not send PHI to public test sandboxes",
    ]:
        add_bullet(doc, item)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading("Utility — FHIR Query Validator", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Version 0.1.0  |  Delivery: May 15, 2026")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 1. Introduction
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "The FHIR Query Validator is a pre-flight utility for FHIR REST search queries. "
        "Given a URL such as Patient?gender=male, it checks syntax and semantics against "
        "a FHIR server's declared capabilities before the search request is sent."
    )
    doc.add_paragraph(
        "It is implemented as a Python library and command-line tool (fhir-validate), "
        "designed for integration into data pipelines, backend services, CI workflows, "
        "and exploratory notebooks."
    )

    # 2. Need for it
    doc.add_heading("2. Need for It", level=1)
    doc.add_paragraph("Healthcare systems query FHIR servers using REST search URLs, but failures are common:")
    for item in [
        "Server heterogeneity — each server supports different resource types, search parameters, modifiers, and comparators",
        "Late failure — invalid queries are often discovered only at runtime after network round-trips",
        "Weak pre-flight checks — syntactically valid URLs can still reference unsupported parameters or invalid coded values",
        "Notebook-to-production gap — validation logic in notebooks is hard to reuse in automation without a structured package",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph(
        "Affected personas include data engineers, backend developers, clinical informaticists, "
        "and QA engineers who need fast, deterministic validation before hitting live servers."
    )

    # 3. What does it do
    doc.add_heading("3. What Does It Do?", level=1)
    doc.add_paragraph("The validator:")
    for item in [
        "Parses a FHIR search URL into resource type and query parameters",
        "Fetches the target server's CapabilityStatement from /metadata",
        "Validates the query against server-declared search parameters, modifiers, and comparators",
        "Applies supplemental static value-set and identifier rules where defined",
        "Returns a structured result: {valid: true/false, errors: [list of messages]}",
    ]:
        add_numbered(doc, item)

    doc.add_paragraph("Example:")
    add_code_block(
        doc,
        'fhir-validate "https://hapi.fhir.org/baseR4/Patient?gender=male"\n'
        "Valid: True",
    )

    # 4. Functionality
    doc.add_heading("4. Functionality", level=1)
    doc.add_paragraph("Validation is server-aware: rules are derived from each server's CapabilityStatement, not a static global registry.")
    validations = [
        ("Resource type", "Must be supported by the server"),
        ("Search parameters", "Must exist for that resource type"),
        ("Modifiers / comparators", "e.g. :exact, :gt must be allowed for that param"),
        ("Static value sets", "Known coded values (e.g. Patient.gender)"),
        ("Patient identifiers", "Numeric, 8–10 digits, not all identical"),
    ]
    table = doc.add_table(rows=1 + len(validations), cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Check"
    table.rows[0].cells[1].text = "Description"
    set_cell_shading(table.rows[0].cells[0], "D5E8F0")
    set_cell_shading(table.rows[0].cells[1], "D5E8F0")
    for i, (check, desc) in enumerate(validations, 1):
        table.rows[i].cells[0].text = check
        table.rows[i].cells[1].text = desc

    add_architecture_table(doc)

    # 5. Repo structure
    add_repo_structure_table(doc)

    # 6. Out of scope
    add_out_of_scope(doc)

    # 7. Tests
    add_testing_section(doc)

    # Appendix
    doc.add_page_break()
    add_appendix_configuration(doc)

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()